from flask import Flask, request, send_file, render_template_string
import os, uuid, docx_to_qti, tempfile, zipfile, io
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

app = Flask(__name__)
app.secret_key = 'qti_final_verified_2026_full'
STORAGE = Path("/tmp/qti_storage")
STORAGE.mkdir(exist_ok=True)

# --- TUTORIAL CONTENT ---
TUTORIAL_TEMPLATE = r'''
<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><script src="https://cdn.tailwindcss.com"></script></head>
<body class="bg-gray-100 min-h-screen p-10 font-sans">
    <div class="max-w-4xl mx-auto bg-white rounded-[2rem] shadow-2xl p-12 border border-gray-100">
        <a href="/" class="text-blue-600 font-bold mb-6 block underline uppercase tracking-widest text-xs">← Back to Converter</a>
        <h1 class="text-4xl font-black mb-10 uppercase tracking-tight text-gray-900">Safe Migration Guide</h1>
        <div class="space-y-12">
            <section class="bg-blue-50 p-8 rounded-3xl border border-blue-100">
                <h2 class="text-2xl font-black text-blue-700 mb-4 uppercase italic">1. Import as Classic Quiz</h2>
                <ol class="space-y-4 list-decimal list-inside text-gray-700 font-medium text-lg">
                    <li>Download your ZIPs from the results page.</li>
                    <li>In Canvas, go to <strong>Settings</strong> > <strong>Import Course Content</strong>.</li>
                    <li>Select <strong>QTI .zip file</strong>.</li>
                    <li><strong>Important:</strong> Do <span class="text-red-600 underline font-bold uppercase">NOT</span> check "Import as New Quiz".</li>
                </ol>
            </section>
            <section class="bg-green-50 p-8 rounded-3xl border border-green-100">
                <h2 class="text-2xl font-black text-green-700 mb-4 uppercase italic">2. Use the Migrate Button</h2>
                <ol class="space-y-4 list-decimal list-inside text-gray-700 font-medium text-lg">
                    <li>Go to the <strong>Quizzes</strong> tab in Canvas.</li>
                    <li>Locate your imported quiz (Classic Rocket icon).</li>
                    <li>Click the <strong>three dots (⋮)</strong> and select <strong>Migrate</strong>.</li>
                </ol>
            </section>
        </div>
    </div>
</body>
</html>
'''

# --- MAIN UI ---
HTML_TEMPLATE = r'''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>QTI Converter Pro</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        .template-box { @apply flex flex-col items-center justify-center p-6 border-2 border-gray-100 rounded-3xl hover:border-blue-500 hover:bg-blue-50 transition cursor-pointer text-center h-full; }
        .step-header { @apply text-xs font-black uppercase tracking-widest text-blue-600 mb-2; }
    </style>
</head>
<body class="bg-gray-100 min-h-screen font-sans flex flex-col">
    <nav class="bg-white border-b border-gray-200 py-6 px-8 mb-8 shadow-sm">
        <div class="max-w-7xl mx-auto flex justify-between items-center">
            <h1 class="text-3xl font-black text-gray-900 tracking-tighter uppercase italic">QTI Converter</h1>
            <a href="/tutorial" class="px-6 py-2 bg-gray-900 text-white text-xs font-black rounded-full hover:bg-blue-600 transition uppercase tracking-widest">Tutorial</a>
        </div>
    </nav>

    <main class="max-w-7xl mx-auto px-4 flex-grow">
        <div class="grid grid-cols-1 md:grid-cols-2 gap-8 items-stretch">
            <div class="flex flex-col">
                <div class="bg-white p-10 rounded-[3rem] shadow-xl border border-gray-50 flex-grow">
                    <p class="step-header">Step 1</p>
                    <h2 class="text-3xl font-black text-gray-900 mb-8 uppercase tracking-tighter">Get a Template</h2>
                    <div class="grid grid-cols-2 gap-6">
                        <a href="/download-template/txt" class="template-box">
                            <span class="text-3xl mb-2">📄</span><span class="font-bold text-gray-800 uppercase tracking-tight">TXT</span>
                        </a>
                        <a href="/download-template/md" class="template-box">
                            <span class="text-3xl mb-2">Ⓜ️</span><span class="font-bold text-gray-800 uppercase tracking-tight">MD</span>
                        </a>
                        <a href="/download-template/csv" class="template-box">
                            <span class="text-3xl mb-2">📊</span><span class="font-bold text-gray-800 uppercase tracking-tight">CSV</span>
                        </a>
                        <a href="/download-template/docx" class="template-box">
                            <span class="text-3xl mb-2">📝</span><span class="font-bold text-gray-800 uppercase tracking-tight">DOCX</span>
                        </a>
                    </div>
                </div>
            </div>

            <div class="flex flex-col">
                <div class="bg-white p-10 rounded-[3rem] shadow-xl border border-gray-100 flex-grow">
                    <p class="step-header">Step 2</p>
                    <h2 class="text-3xl font-black text-gray-900 mb-8 uppercase tracking-tighter">Convert Quiz</h2>
                    {% if results %}
                        <div class="space-y-4">
                            {% for r in results %}
                            <div class="flex items-center justify-between bg-blue-50 p-5 rounded-3xl border border-blue-100">
                                <span class="font-black text-blue-900 truncate mr-4 text-sm">{{ r.name }}</span>
                                <a href="/download/{{ r.id }}/{{ r.name }}" class="bg-blue-600 text-white font-black px-6 py-3 rounded-2xl text-xs uppercase tracking-widest hover:bg-blue-700 transition">Download</a>
                            </div>
                            {% endfor %}
                            <a href="/" class="block text-center pt-6 text-blue-600 font-bold uppercase text-xs tracking-widest underline underline-offset-4 tracking-[0.2em] hover:text-blue-800 transition">← Convert More</a>
                        </div>
                    {% else %}
                        <form action="/" method="post" enctype="multipart/form-data" id="upload-form" class="h-full flex flex-col">
                            <div id="drop-zone" class="border-4 border-dashed border-gray-100 rounded-[2.5rem] flex-grow flex flex-col items-center justify-center hover:bg-gray-50 hover:border-blue-400 transition cursor-pointer min-h-[300px]">
                                <input type="file" name="files" id="file-input" multiple class="hidden" accept=".docx,.md,.txt,.csv">
                                <div class="bg-blue-100 p-4 rounded-full mb-4">
                                    <svg class="w-10 h-10 text-blue-600" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M7 16a4 4 0 01-.88-7.903A5 5 0 1115.9 6L16 6a5 5 0 011 9.9M15 13l-3-3m0 0l-3 3m3-3v12"></path></svg>
                                </div>
                                <p class="text-gray-400 font-black uppercase tracking-widest text-sm text-center px-4">Drag Files Here or Click to Browse</p>
                            </div>
                            <div id="file-list" class="mt-6 space-y-2"></div>
                            <button type="button" id="submit-btn" class="w-full bg-blue-600 text-white font-black py-5 rounded-[2rem] shadow-2xl mt-6 hidden uppercase tracking-widest text-lg hover:bg-blue-700 transition">Process Files</button>
                        </form>
                    {% endif %}
                </div>
            </div>
        </div>
    </main>
    <footer class="py-12 text-center text-gray-400 text-[10px] font-black uppercase tracking-[0.3em]">
        <p>&copy; 2026 Eric Dannewitz & Gemini &bull; All Rights Reserved</p>
    </footer>
    <script>
        const dropZone = document.getElementById('drop-zone');
        const fileInput = document.getElementById('file-input');
        const fileList = document.getElementById('file-list');
        const submitBtn = document.getElementById('submit-btn');
        const form = document.getElementById('upload-form');
        let currentFiles = [];

        if (dropZone) {
            dropZone.onclick = () => fileInput.click();
            fileInput.onchange = (e) => handleFiles(Array.from(e.target.files));
            dropZone.ondragover = (e) => { e.preventDefault(); dropZone.classList.add('bg-blue-50'); };
            dropZone.ondragleave = () => dropZone.classList.remove('bg-blue-50');
            dropZone.ondrop = (e) => { e.preventDefault(); dropZone.classList.remove('bg-blue-50'); handleFiles(Array.from(e.dataTransfer.files)); };
        }

        function handleFiles(files) {
            currentFiles = files;
            const dt = new DataTransfer();
            currentFiles.forEach(f => dt.items.add(f));
            fileInput.files = dt.files;
            fileList.innerHTML = currentFiles.map(f => `<div class="px-4 py-2 bg-gray-50 border-2 border-gray-100 rounded-xl text-[10px] font-black uppercase text-gray-400 tracking-wider">${f.name}</div>`).join('');
            submitBtn.classList.toggle('hidden', currentFiles.length === 0);
        }

        if (submitBtn) {
            submitBtn.onclick = () => {
                if (fileInput.files.length > 0) {
                    submitBtn.innerText = "Processing...";
                    submitBtn.disabled = true;
                    form.submit();
                }
            };
        }
    </script>
</body>
</html>
'''

# --- ROUTES ---
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        job_id = str(uuid.uuid4()); job_dir = STORAGE / job_id; job_dir.mkdir()
        results = []
        for f in request.files.getlist('files'):
            if f.filename == '': continue
            name = Path(f.filename).stem; input_path = job_dir / f.filename; f.save(str(input_path))
            zip_name = f"{name}_qti.zip"; output_zip = job_dir / zip_name
            if docx_to_qti.convert_to_qti_zip(input_path, output_zip):
                results.append({'name': zip_name, 'id': job_id})
        return render_template_string(HTML_TEMPLATE, results=results)
    return render_template_string(HTML_TEMPLATE, results=None)

@app.route('/download-template/<fmt>')
def download_template(fmt):
    buf = io.BytesIO()
    content = r'''Title: Q1 Multiple Choice
Points: 1
What is the primary color of a clear sky?
*a) Blue
b) Green

Title: Q2 True or False
Points: 1
The Earth is flat.
*a) FALSE
b) TRUE

Title: Q3 Fill-in-the-blank
Points: 1
Plants make food via [photosynthesis].

Title: Q4 Multiple Answers
Points: 1
Which are prime?
[x] 2
[x] 3
[ ] 4
[x] 5

Title: Q5 Matching
Points: 1
Match capitals:
UK -> London
Japan -> Tokyo

Title: Q6 Numerical Answer
Points: 1
Value of Pi (2 decimals)?
[answer] 3.14

Title: Q7 Formula
Points: 1
Calculate [speed] * [time]
[formula] speed * time

Title: Q8 Essay
Points: 1
Explain the industrial revolution.

Title: Q9 File Upload
Points: 1
Upload your report.
'''

    if fmt in ['txt', 'md']:
        buf.write(content.replace(r'\n', '\n').encode('utf-8')); name = f"Template.{fmt}"
    elif fmt == 'csv':
        # FIXED CSV Output to include all 9 question types with 4 choice columns
        csv_data = (
            b"Type,Title,Points,Question,Answer,Choice1,Choice2,Choice3,Choice4\n"
            b"MC,Q1 Multiple Choice,1,What is the primary color of a clear sky?,Blue,Blue,Green,,\n"
            b"TF,Q2 True or False,1,The Earth is flat.,FALSE,TRUE,FALSE,,\n"
            b"FIB,Q3 Fill-in-the-blank,1,Plants make food via [photosynthesis].,photosynthesis,,,,\n"
            b"MR,Q4 Multiple Answers,1,Which are prime?,124,2,3,4,5\n"
            b"Matching,Q5 Matching,1,Match capitals:,,UK->London,Japan->Tokyo,,\n"
            b"Numeric,Q6 Numerical Answer,1,Value of Pi (2 decimals)?,3.14,,,,\n"
            b"Formula,Q7 Formula,1,Calculate [speed] * [time],,,,,\n"
            b"Essay,Q8 Essay,1,Explain the industrial revolution.,,,,,\n"
            b"FileUpload,Q9 File Upload,1,Upload your report.,,,,,\n"
        )
        buf.write(csv_data); name = "Template.csv"
    elif fmt == 'docx':
        doc = Document(); doc.add_heading('Comprehensive Quiz Template', 0)
        for block in content.split('Title:'):
            if not block.strip(): continue
            lines = block.strip().split(r'\n')
            doc.add_paragraph(f"Title: {lines[0]}")
            for line in lines[1:]:
                p = doc.add_paragraph(line)
                if line.startswith('*') or '[x]' in line:
                    p.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        doc.save(buf); name = "Template.docx"

    buf.seek(0)
    return send_file(buf, download_name=name, as_attachment=True)

@app.route('/tutorial')
def tutorial(): return render_template_string(TUTORIAL_TEMPLATE)

@app.route('/download/<job_id>/<filename>')
def download(job_id, filename): return send_file(STORAGE / job_id / filename, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
